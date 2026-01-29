# document_chat.py - Chat with Uploaded Documents Feature
# This module provides functionality for users to upload documents and chat about them

import os
import re
import logging
import tempfile
import hashlib
from typing import List, Dict, Any, Optional
from io import BytesIO

# PDF and Document Processing
from PyPDF2 import PdfReader

# Image Processing (for scanned documents)
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# DOCX Processing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# LangChain for text splitting
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Environment
from dotenv import load_dotenv

load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt', 'png', 'jpg', 'jpeg'}


class DocumentProcessor:
    """Process uploaded documents and extract text"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            
            full_text = "\n\n".join(text_parts)
            
            # If no text extracted (scanned PDF), try OCR
            if not full_text.strip() and OCR_AVAILABLE:
                logger.info("No text found in PDF, attempting OCR...")
                full_text = self._ocr_pdf(file_content)
            
            return full_text
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return ""
    
    def _ocr_pdf(self, file_content: bytes) -> str:
        """OCR scanned PDF using pytesseract"""
        try:
            import pypdfium2 as pdfium
            
            pdf = pdfium.PdfDocument(file_content)
            text_parts = []
            
            for page_num in range(len(pdf)):
                page = pdf[page_num]
                # Render page at 300 DPI for better OCR
                bitmap = page.render(scale=300/72)
                pil_image = bitmap.to_pil()
                
                # OCR the image
                page_text = pytesseract.image_to_string(pil_image)
                if page_text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"OCR error: {e}")
            return ""
    
    def extract_text_from_image(self, file_content: bytes) -> str:
        """Extract text from image using OCR"""
        if not OCR_AVAILABLE:
            logger.warning("OCR not available - pytesseract not installed")
            return ""
        
        try:
            image = Image.open(BytesIO(file_content))
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            logger.error(f"Image OCR error: {e}")
            return ""
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            logger.warning("DOCX processing not available - python-docx not installed")
            return ""
        
        try:
            doc = DocxDocument(BytesIO(file_content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ""
    
    def extract_text_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try UTF-8 first, then fallback to other encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            return ""
    
    def process_document(self, filename: str, file_content: bytes) -> Dict[str, Any]:
        """Process document and return extracted text and metadata"""
        ext = filename.lower().split('.')[-1]
        
        if ext not in ALLOWED_EXTENSIONS:
            return {"error": f"Unsupported file type: {ext}", "text": "", "filename": filename}
        
        # Extract text based on file type
        if ext == 'pdf':
            text = self.extract_text_from_pdf(file_content)
        elif ext in ['png', 'jpg', 'jpeg']:
            text = self.extract_text_from_image(file_content)
        elif ext in ['docx', 'doc']:
            text = self.extract_text_from_docx(file_content)
        elif ext == 'txt':
            text = self.extract_text_from_txt(file_content)
        else:
            text = ""
        
        # Generate document hash for caching
        doc_hash = hashlib.md5(file_content).hexdigest()
        
        # Split into chunks for better retrieval
        chunks = self.text_splitter.split_text(text) if text else []
        
        return {
            "filename": filename,
            "text": text,
            "chunks": chunks,
            "hash": doc_hash,
            "char_count": len(text),
            "chunk_count": len(chunks),
            "file_type": ext
        }


class DocumentChatSession:
    """Manages a chat session with an uploaded document"""
    
    def __init__(self, document_data: Dict[str, Any], llm):
        self.document_data = document_data
        self.llm = llm
        self.chat_history = []
        
        # Create the document context prompt
        self.document_context = self._build_context()
    
    def _build_context(self) -> str:
        """Build context from document text"""
        # Use the full text if it's not too long, otherwise use chunks
        text = self.document_data.get("text", "")
        
        # Limit context to avoid token overflow (roughly 4000 tokens ~ 16000 chars)
        max_context_chars = 16000
        if len(text) > max_context_chars:
            # Use first and last parts
            half = max_context_chars // 2
            text = text[:half] + "\n\n[... document continues ...]\n\n" + text[-half:]
        
        return text
    
    def get_relevant_context(self, query: str) -> str:
        """Get relevant chunks based on query (simple keyword matching)"""
        query_words = set(query.lower().split())
        chunks = self.document_data.get("chunks", [])
        
        if not chunks:
            return self.document_context
        
        # Score each chunk by keyword overlap
        scored_chunks = []
        for chunk in chunks:
            chunk_words = set(chunk.lower().split())
            overlap = len(query_words.intersection(chunk_words))
            scored_chunks.append((overlap, chunk))
        
        # Sort by relevance and take top chunks
        scored_chunks.sort(key=lambda x: x[0], reverse=True)
        top_chunks = [chunk for score, chunk in scored_chunks[:5] if score > 0]
        
        if top_chunks:
            return "\n\n---\n\n".join(top_chunks)
        
        # If no relevant chunks found, return first few chunks
        return "\n\n---\n\n".join(chunks[:3]) if chunks else self.document_context
    
    def chat(self, user_query: str, legal_knowledge_response: Optional[str] = None) -> Dict[str, Any]:
        """
        Process a chat query about the document
        
        Args:
            user_query: The user's question
            legal_knowledge_response: Optional response from the main RAG system (legal knowledge)
        """
        try:
            # Get relevant context from the uploaded document
            doc_context = self.get_relevant_context(user_query)
            
            # Build the prompt
            prompt = self._build_chat_prompt(user_query, doc_context, legal_knowledge_response)
            
            # Get response from LLM
            response = self.llm.invoke(prompt)
            answer = response.content if hasattr(response, 'content') else str(response)
            
            # Store in history
            self.chat_history.append({
                "role": "user",
                "content": user_query
            })
            self.chat_history.append({
                "role": "assistant", 
                "content": answer
            })
            
            return {
                "answer": answer,
                "document_name": self.document_data.get("filename", "Unknown"),
                "sources": [{
                    "type": "uploaded_document",
                    "filename": self.document_data.get("filename")
                }]
            }
            
        except Exception as e:
            logger.error(f"Document chat error: {e}")
            return {
                "answer": f"I encountered an error processing your question: {str(e)}",
                "error": True
            }
    
    def _build_chat_prompt(self, query: str, doc_context: str, legal_knowledge: Optional[str] = None) -> str:
        """Build the chat prompt with document context and optional legal knowledge"""
        
        prompt = f"""You are a legal assistant helping analyze an uploaded document.

DOCUMENT CONTEXT (from the user's uploaded document "{self.document_data.get('filename', 'document')}"):
---
{doc_context}
---
"""
        
        if legal_knowledge:
            prompt += f"""
RELEVANT LEGAL KNOWLEDGE (from legal database):
---
{legal_knowledge}
---
"""
        
        prompt += f"""
USER QUESTION: {query}

INSTRUCTIONS:
1. Answer based PRIMARILY on the uploaded document content above
2. If legal knowledge is provided, use it to give legal context and implications
3. Be specific - cite clauses, sections, or page numbers from the document when possible
4. If the document doesn't contain information to answer the question, say so clearly
5. Be helpful and explain legal terms in simple language

YOUR ANSWER:"""
        
        return prompt


class DocumentChatManager:
    """Manages document chat sessions with LLM integration"""
    
    def __init__(self, chatbot_manager=None):
        """
        Initialize with optional chatbot_manager for legal knowledge integration
        
        Args:
            chatbot_manager: The main ChatbotManager instance for accessing legal knowledge
        """
        self.processor = DocumentProcessor()
        self.chatbot_manager = chatbot_manager
        self.active_sessions: Dict[str, DocumentChatSession] = {}
        
        # Initialize LLM (same as main chatbot)
        self.llm = self._initialize_llm()
    
    def _initialize_llm(self):
        """Initialize the LLM based on environment"""
        app_env = os.getenv('APP_ENV', 'local')
        
        if app_env == 'production':
            from langchain_groq import ChatGroq
            groq_api_key = os.getenv("GROQ_API_KEY")
            if not groq_api_key:
                raise ValueError("GROQ_API_KEY not found")
            
            return ChatGroq(
                model_name="llama-3.3-70b-versatile",
                temperature=0.3,
                groq_api_key=groq_api_key
            )
        else:
            from langchain_ollama import ChatOllama
            base_url = os.getenv('OLLAMA_BASE_URL', 'http://192.168.0.25:11434')
            model = os.getenv('OLLAMA_MODEL', 'qwen2.5:14b')
            
            return ChatOllama(
                model=model,
                temperature=0.3,
                base_url=base_url,
                timeout=120
            )
    
    def upload_document(self, filename: str, file_content: bytes, session_id: str) -> Dict[str, Any]:
        """
        Upload and process a document for chat
        
        Args:
            filename: Name of the uploaded file
            file_content: Raw bytes of the file
            session_id: Unique session identifier
        
        Returns:
            Dict with processing result
        """
        # Validate file size
        if len(file_content) > MAX_FILE_SIZE:
            return {
                "success": False,
                "error": f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
            }
        
        # Process the document
        doc_data = self.processor.process_document(filename, file_content)
        
        if "error" in doc_data:
            return {
                "success": False,
                "error": doc_data["error"]
            }
        
        if not doc_data.get("text"):
            return {
                "success": False,
                "error": "Could not extract text from document. Please try a different file."
            }
        
        # Create chat session
        chat_session = DocumentChatSession(doc_data, self.llm)
        self.active_sessions[session_id] = chat_session
        
        return {
            "success": True,
            "session_id": session_id,
            "filename": filename,
            "char_count": doc_data["char_count"],
            "chunk_count": doc_data["chunk_count"],
            "file_type": doc_data["file_type"],
            "preview": doc_data["text"][:500] + "..." if len(doc_data["text"]) > 500 else doc_data["text"]
        }
    
    def chat_with_document(
        self, 
        session_id: str, 
        query: str,
        include_legal_knowledge: bool = True
    ) -> Dict[str, Any]:
        """
        Chat about an uploaded document
        
        Args:
            session_id: The document session ID
            query: User's question
            include_legal_knowledge: Whether to include responses from legal knowledge base
        
        Returns:
            Dict with answer and metadata
        """
        if session_id not in self.active_sessions:
            return {
                "success": False,
                "error": "No document uploaded for this session. Please upload a document first."
            }
        
        session = self.active_sessions[session_id]
        
        # Optionally get legal knowledge context
        legal_knowledge = None
        if include_legal_knowledge and self.chatbot_manager:
            try:
                # Get relevant legal knowledge from the main RAG system
                legal_response = self.chatbot_manager.get_response(
                    query, 
                    use_rag=True,
                    enable_content_filter=False,
                    enable_pii_detection=False
                )
                legal_knowledge = legal_response.get("answer", "")
            except Exception as e:
                logger.warning(f"Could not get legal knowledge: {e}")
        
        # Chat with the document
        result = session.chat(query, legal_knowledge)
        result["success"] = True
        result["session_id"] = session_id
        
        return result
    
    def get_session_info(self, session_id: str) -> Optional[Dict[str, Any]]:
        """Get information about a document chat session"""
        if session_id not in self.active_sessions:
            return None
        
        session = self.active_sessions[session_id]
        return {
            "filename": session.document_data.get("filename"),
            "char_count": session.document_data.get("char_count"),
            "chat_history_length": len(session.chat_history)
        }
    
    def clear_session(self, session_id: str) -> bool:
        """Clear a document chat session"""
        if session_id in self.active_sessions:
            del self.active_sessions[session_id]
            return True
        return False


# Singleton instance
_document_chat_manager: Optional[DocumentChatManager] = None

def get_document_chat_manager(chatbot_manager=None) -> DocumentChatManager:
    """Get or create the document chat manager singleton"""
    global _document_chat_manager
    if _document_chat_manager is None:
        _document_chat_manager = DocumentChatManager(chatbot_manager)
    return _document_chat_manager
